#!/usr/bin/env python3
"""Build shareable exports of DESIGN.md.

    python3 docs/design/build_exports.py

Produces, next to DESIGN.md:
  "CalBlue Platform Design (draft v0.3).docx"
                 Import into any Google account (Docs > File > Open > Upload).
                 This is the reliable path: Docs' .docx importer keeps the
                 diagrams and tables intact, and names the new doc after the
                 file — hence the deliberately human filename.
  DESIGN.html    Same content, diagrams inlined as base64. A fallback, and
                 openable in a browser without Word or Docs.

Both are generated from DESIGN.md and diagrams/*.png, so DESIGN.md stays the
single source of truth. Re-run after editing either.
"""

import base64
import html as H
import io
import os
import re

from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "DESIGN.md")

# heading prefix -> diagram rendered underneath it
DIAGRAMS = {
    "Diagram 1": "diagrams/02-permissions.png",
    "Diagram 2": "diagrams/01-entity-map.png",
    "Diagram 3": "diagrams/03-lifecycles.png",
    "Diagram 4": "diagrams/04-flow.png",
    "Diagram 5": "diagrams/05-billing.png",
    "Diagram 6": "diagrams/06-hosting.png",
    "Diagram 7": "diagrams/07-clients.png",
}

TITLE = ("CalBlue — League, Pickup & Membership Platform: "
         "Data Model & Permissions Design (draft v0.3)")

# Google Docs names an imported document after its filename, so this is what
# the club will see in their Drive.
DOCX_NAME = "CalBlue Platform Design (draft v0.3).docx"

TOKEN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def tokenize(text):
    """Split inline markdown into (text, bold, italic, code) runs."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            yield part[2:-2], True, False, False
        elif part.startswith("`") and part.endswith("`"):
            yield part[1:-1], False, False, True
        elif part.startswith("*") and part.endswith("*"):
            yield part[1:-1], False, True, False
        else:
            yield part, False, False, False


# --------------------------------------------------------------- parse
def parse(md):
    """DESIGN.md -> a flat list of blocks the writers below can render."""
    blocks, lines, i = [], md.split("\n"), 0
    while i < len(lines):
        ln = lines[i]

        m = re.match(r"^(#{1,4}) (.*)", ln)
        if m:
            level, text = len(m.group(1)), m.group(2)
            blocks.append(("heading", level, text))
            key = text.split(" —")[0].strip()
            if key in DIAGRAMS:
                blocks.append(("image", DIAGRAMS[key]))
            i += 1
            continue

        if ln.strip() == "---":
            blocks.append(("rule",))
            i += 1
            continue

        if ln.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            rows = [r for r in rows
                    if not all(set(c) <= set("-: ") for c in r)]
            blocks.append(("table", rows))
            continue

        ordered = re.match(r"^\d+\. ", ln)
        if ordered or ln.startswith("- "):
            items = []
            while i < len(lines) and (re.match(r"^\d+\. ", lines[i])
                                      or lines[i].startswith("- ")
                                      or lines[i].startswith("  ")):
                item = re.sub(r"^(\d+\. |- )", "", lines[i]).strip()
                if lines[i].startswith("  ") and items:
                    items[-1] += " " + item      # continuation line
                else:
                    items.append(item)
                i += 1
            blocks.append(("list", bool(ordered), items))
            continue

        if ln.strip():
            para = [ln]
            i += 1
            while (i < len(lines) and lines[i].strip()
                   and not re.match(r"^([#\-|]|\d+\. )", lines[i])):
                para.append(lines[i])
                i += 1
            blocks.append(("para", " ".join(para)))
            continue

        i += 1
    return blocks


# ---------------------------------------------------------------- docx
def write_docx(blocks, out):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.8)
        section.top_margin = section.bottom_margin = Inches(0.8)
    usable = doc.sections[0].page_width - (doc.sections[0].left_margin
                                           + doc.sections[0].right_margin)

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica Neue"
    normal.font.size = Pt(10.5)

    def runs(paragraph, text, italic=False):
        for chunk, bold, ital, code in tokenize(text):
            run = paragraph.add_run(chunk)
            run.bold = bold
            run.italic = ital or italic
            if code:
                run.font.name = "Menlo"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    for block in blocks:
        kind = block[0]

        if kind == "heading":
            _, level, text = block
            para = doc.add_heading(level=min(level, 4))
            runs(para, text)

        elif kind == "para":
            text = block[1]
            para = doc.add_paragraph()
            # a whole-line *italic* is a diagram caption
            if text.startswith("*") and text.endswith("*") and "**" not in text:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                runs(para, text[1:-1], italic=True)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x5B, 0x6B, 0x83)
            else:
                runs(para, text)

        elif kind == "list":
            _, ordered, items = block
            for item in items:
                para = doc.add_paragraph(
                    style="List Number" if ordered else "List Bullet")
                runs(para, item)

        elif kind == "table":
            rows = block[1]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    para = table.cell(r, c).paragraphs[0]
                    runs(para, cell)
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if r == 0:
                            run.bold = True

        elif kind == "image":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(os.path.join(HERE, block[1]),
                                       width=usable)

        elif kind == "rule":
            doc.add_paragraph()

    doc.core_properties.title = TITLE
    doc.save(out)
    return out


# ---------------------------------------------------------------- html
def embed(path, width=1700):
    im = Image.open(path)
    if im.width > width:
        im = im.resize((width, round(im.height * width / im.width)),
                       Image.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, "PNG", optimize=True)
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()


def inline_html(text):
    out = []
    for chunk, bold, ital, code in tokenize(text):
        chunk = H.escape(chunk)
        if bold:
            chunk = f"<strong>{chunk}</strong>"
        elif ital:
            chunk = f"<em>{chunk}</em>"
        elif code:
            chunk = f"<code>{chunk}</code>"
        out.append(chunk)
    return "".join(out)


def write_html(blocks, out):
    body = []
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            body.append(f"<h{level}>{inline_html(text)}</h{level}>")
        elif kind == "para":
            body.append(f"<p>{inline_html(block[1])}</p>")
        elif kind == "rule":
            body.append("<hr>")
        elif kind == "list":
            _, ordered, items = block
            tag = "ol" if ordered else "ul"
            body.append(f"<{tag}>"
                        + "".join(f"<li>{inline_html(x)}</li>" for x in items)
                        + f"</{tag}>")
        elif kind == "table":
            rows = block[1]
            cells = ["<tr>" + "".join(f"<th>{inline_html(c)}</th>"
                                      for c in rows[0]) + "</tr>"]
            cells += ["<tr>" + "".join(f"<td>{inline_html(c)}</td>"
                                       for c in r) + "</tr>" for r in rows[1:]]
            body.append('<table border="1" cellpadding="6" '
                        'style="border-collapse:collapse">'
                        + "".join(cells) + "</table>")
        elif kind == "image":
            src = embed(os.path.join(HERE, block[1]))
            body.append(f'<p><img src="{src}" style="width:100%"></p>')

    doc = (f"<html><head><meta charset='utf-8'><title>{H.escape(TITLE)}</title>"
           "</head><body style=\"font-family:Helvetica,Arial,sans-serif;"
           "line-height:1.5;max-width:60em\">"
           + "\n".join(body) + "</body></html>")
    with open(out, "w") as fh:
        fh.write(doc)
    return out


if __name__ == "__main__":
    blocks = parse(open(MD).read())
    for path in (write_docx(blocks, os.path.join(HERE, DOCX_NAME)),
                 write_html(blocks, os.path.join(HERE, "DESIGN.html"))):
        print(f"{path}  ({os.path.getsize(path) / 1_000_000:.2f} MB)")
